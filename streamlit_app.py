import streamlit as st
from openai import OpenAI
from docx import Document

# Función para extraer texto de un archivo Word (incluyendo tablas)
def extract_text_from_docx(file):
    doc = Document(file)
    full_text = []
    
    # Extraer párrafos
    for para in doc.paragraphs:
        full_text.append(para.text)
        
    # Extraer texto de tablas (importante para reportes)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))
            
    return "\n".join(full_text)

# Título y configuración
st.title("🤖 Chatbot Documental")
st.write(
    "Sube tu documento Word (.docx) y haz preguntas sobre su contenido. "
    "Este bot utilizará la información del archivo para responderte."
)

# Input para la API Key
openai_api_key = st.text_input("OpenAI API Key", type="password")

# 1. Componente para subir el archivo
uploaded_file = st.file_uploader("Sube tu documento Word aquí", type=["docx"])

# Variable para almacenar el contenido del documento en la sesión
if "document_context" not in st.session_state:
    st.session_state.document_context = ""

# Procesar el archivo si se sube uno nuevo
if uploaded_file is not None:
    try:
        # Extraer el texto y guardarlo en el estado
        text_content = extract_text_from_docx(uploaded_file)
        st.session_state.document_context = text_content
        st.success(f"✅ Documento procesado con éxito. ¡Ya puedes preguntar!")
        
        # Opcional: Mostrar un poco del texto extraído para verificar (oculto en un expander)
        with st.expander("Ver contenido extraído"):
            st.write(text_content[:1000] + "...") # Muestra los primeros 1000 caracteres
            
    except Exception as e:
        st.error(f"Error al leer el archivo: {e}")

# Verificación de API Key
if not openai_api_key:
    st.info("Por favor ingresa tu OpenAI API Key para continuar.", icon="🗝️")
else:
    client = OpenAI(api_key=openai_api_key)

    # Inicializar historial de chat
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Mostrar mensajes previos
    for message in st.session_state.messages:
        # No mostramos el mensaje del sistema (el contexto) al usuario final
        if message["role"] != "system":
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

    # Input del usuario
    if prompt := st.chat_input("Pregunta algo sobre el documento..."):

        # Guardar y mostrar mensaje del usuario
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        # 2. Construir la lista de mensajes con el contexto del documento
        # Si hay un documento cargado, lo añadimos como instrucción del sistema
        messages_to_send = []
        
        if st.session_state.document_context:
            system_prompt = {
                "role": "system", 
                "content": f"Eres un asistente útil. Responde a las preguntas del usuario basándote ÚNICAMENTE en el siguiente contexto extraído de un documento:\n\n{st.session_state.document_context}"
            }
            messages_to_send.append(system_prompt)
        else:
            # Si no hay doc, actúa como un bot normal
            messages_to_send.append({"role": "system", "content": "Eres un asistente útil."})

        # Añadir el historial de la conversación (para que tenga memoria)
        messages_to_send.extend(st.session_state.messages)

        # Generar respuesta
        stream = client.chat.completions.create(
            model="gpt-4o-mini", # Usamos gpt-4o-mini porque soporta textos mucho más largos que 3.5
            messages=messages_to_send,
            stream=True,
        )

        # Mostrar respuesta
        with st.chat_message("assistant"):
            response = st.write_stream(stream)
        
        st.session_state.messages.append({"role": "assistant", "content": response})
